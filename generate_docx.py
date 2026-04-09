"""
Generate DMS_Thesis_Theory_Handbook.docx - properly formatted.
Uses python-docx library.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ---- Style customization ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15

# Heading styles
HEADING_COLORS = {
    1: ('Calibri', 26, '1B3A5C', 24, 12),
    2: ('Calibri', 18, '2E5D8C', 16, 8),
    3: ('Calibri', 14, '3A7CA5', 12, 6),
}
for level, (fname, fsize, color_hex, sp_before, sp_after) in HEADING_COLORS.items():
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = fname
    h_style.font.size = Pt(fsize)
    h_style.font.bold = True
    h_style.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    h_style.paragraph_format.space_before = Pt(sp_before)
    h_style.paragraph_format.space_after = Pt(sp_after)
    h_style.paragraph_format.keep_with_next = True

# ---- Helper functions ----

def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_rich_para(parts, align=None, space_after=None):
    """parts = list of (text, bold, italic, size, color, font_name)"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for part in parts:
        text = part[0]
        run = p.add_run(text)
        if len(part) > 1 and part[1]: run.bold = True
        if len(part) > 2 and part[2]: run.italic = True
        if len(part) > 3 and part[3]: run.font.size = Pt(part[3])
        if len(part) > 4 and part[4]: run.font.color.rgb = RGBColor(*bytes.fromhex(part[4]))
        if len(part) > 5 and part[5]: run.font.name = part[5]
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_r = p.add_run(text)
    else:
        # The style already adds the bullet; just set text
        p.text = text
    return p

def add_numbered_item(text):
    p = doc.add_paragraph(style='List Number')
    p.text = text
    return p

def add_code_block(lines):
    """Add a visually distinct code block with gray background."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.0
        # Add shading to paragraph
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # Add light blue background
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.name = 'Cambria Math'

def add_table(headers, rows):
    """Add a well-formatted table."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        # Header shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            # Alternate row shading
            if i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
                tcPr.append(shd)

    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_callout(text, bg_color="FFF3CD", text_color="856404"):
    """Yellow callout box for important notes."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(*bytes.fromhex(text_color))
    run.font.size = Pt(11)

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ==================== DOCUMENT CONTENT ====================

# ---- TITLE PAGE ----
for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

add_para('DYNAMIC MODEL SWITCHING', bold=True, size=32, color='1B3A5C',
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para('FOR REAL-TIME UAV INSPECTION', bold=True, size=32, color='1B3A5C',
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

add_separator()

add_para('Complete Thesis Theory Handbook', italic=True, size=18, color='2E5D8C',
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=24)

# Meta info
meta_items = [
    ('Author:', ' Ganapathi Teja Chirravuri'),
    ('University:', ' Technische Universitat Chemnitz'),
    ('Faculty:', ' Computer Science'),
    ('Supervisor:', ' Dr. Battseren'),
    ('Research Group:', ' Software Architecture for Real-Time Image Analysis'),
    ('Date:', ' April 2026'),
]
for label, value in meta_items:
    add_rich_para([
        (label, True, False, 12, '1B3A5C'),
        (value, False, False, 12, '444444'),
    ], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

doc.add_paragraph()
add_separator()
doc.add_paragraph()

add_para('DMS-Raptor Framework', bold=True, size=13, color='2E5D8C',
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Adaptive per-frame model selection between YOLOv8n (fast) and YOLOv8s (accurate)\n'
         'using scene complexity proxies, confidence-based switching, and actuator stabilization.',
         italic=True, size=10, color='666666',
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('8 switching policies \u2022 12 UAV videos \u2022 78,225 frames \u2022 19 references',
         size=10, color='888888', align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ---- TABLE OF CONTENTS ----
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'Problem Statement & Motivation'),
    ('2.', 'Background & Theoretical Foundations'),
    ('3.', 'Related Work'),
    ('4.', 'System Architecture'),
    ('5.', 'Scene Complexity Proxies'),
    ('6.', 'The 8 Switching Policies'),
    ('7.', 'The C-Score Failure (Central Contribution)'),
    ('8.', 'Actuator Stabilization'),
    ('9.', 'Evaluation Framework'),
    ('10.', 'Experimental Results'),
    ('11.', 'Key Research Findings'),
    ('12.', 'Implementation Details'),
    ('13.', 'Deployment Considerations'),
    ('14.', 'Future Work: Data & Model Agnosticism'),
    ('15.', 'References'),
]
for num, title in toc_items:
    add_rich_para([
        (f'{num}  ', True, False, 12, '1B3A5C'),
        (title, False, False, 12, '333333'),
    ], space_after=4)

doc.add_page_break()

# ==================== CHAPTER 1 ====================
doc.add_heading('1. Problem Statement & Motivation', level=1)

doc.add_heading('1.1 The UAV Inspection Challenge', level=2)
add_para('Unmanned Aerial Vehicles (UAVs) are increasingly deployed for visual inspection of power-line infrastructure, particularly insulators that are critical for electrical grid reliability. These inspections require:')
add_bullet('Real-time processing: Decisions must be made at frame rate to guide flight paths or trigger alerts')
add_bullet('High detection accuracy: Missed defects can lead to grid failures')
add_bullet('Resource-constrained deployment: Edge devices (NVIDIA Jetson Nano, embedded GPUs) have limited compute')

doc.add_heading('1.2 The Speed-Accuracy Dilemma', level=2)
add_para('Modern object detectors like YOLOv8 offer multiple model sizes:')
add_table(
    ['Model', 'Parameters', 'Speed (CPU)', 'Accuracy'],
    [
        ['YOLOv8n (nano)', '3.2M', '~100ms', 'Good (F1=0.870 on glass)'],
        ['YOLOv8s (small)', '11.2M', '~240ms', 'Better (F1=0.892 on glass)'],
    ]
)
add_callout('\u26a1 The latency gap of ~140ms on CPU creates an opportunity for intelligent switching.')

doc.add_heading('1.3 Core Idea', level=2)
add_para('DMS-Raptor is a Dynamic Model Switching framework that analyzes each frame\'s scene complexity and routes it to the appropriate model:', bold=True)
add_bullet('Simple frames (clear sky, distant insulators) \u2192 YOLOv8n (~100ms, good enough)')
add_bullet('Complex frames (cluttered backgrounds, reflections, close-ups) \u2192 YOLOv8s (~240ms, more accurate)')
add_callout('Key Research Question: Can scene complexity proxies reliably predict when the fast model will fail, enabling intelligent per-frame model selection?', bg_color="D6EAF8", text_color="1B4F72")

doc.add_heading('1.4 Thesis Contributions', level=2)
add_numbered_item('Negative result (central contribution): Demonstrating that rolling percentile normalization destroys proxy-target correlation, causing the C-score approach to fail')
add_numbered_item('Reactive switching via conf_ema: A zero-overhead policy using the model\'s own confidence as the switching signal')
add_numbered_item('EMA-relative normalization: An alternative to percentile normalization that preserves signal dynamics (multi_proxy)')
add_numbered_item('Comprehensive evaluation: 8 policies across 12 videos (78,225 frames) with pseudo-oracle validation')
add_numbered_item('Deployment boundary: DMS works on CPU/edge platforms but NOT on high-end GPUs (where n and s have similar latency)')
doc.add_page_break()

# ==================== CHAPTER 2 ====================
doc.add_heading('2. Background & Theoretical Foundations', level=1)

doc.add_heading('2.1 YOLOv8 Architecture', level=2)
add_para('YOLOv8 (Jocher et al., 2023) [1] is a single-stage real-time object detector with:')
add_bullet('CSPDarknet53 backbone: Cross-Stage Partial connections for efficient feature extraction')
add_bullet('PANet neck: Path Aggregation Network for multi-scale feature fusion')
add_bullet('Decoupled head: Separate classification and regression branches')
add_bullet('Anchor-free design: Direct prediction of object centers and sizes')
add_bullet('Built-in NMS: C++/CUDA Non-Maximum Suppression applied internally during .predict()')

add_para('The key architectural difference between n and s variants:', bold=True)
add_bullet('YOLOv8n: Fewer channels per layer (depth_multiple=0.33, width_multiple=0.25)')
add_bullet('YOLOv8s: More channels (depth_multiple=0.33, width_multiple=0.50)')
add_bullet('Both share the same architecture, differing only in capacity')

doc.add_heading('2.2 Image Quality & Complexity Metrics', level=2)

doc.add_heading('Laplacian Variance (L)', level=3)
add_para('The Laplacian operator is a second-order derivative filter that highlights edges:')
add_equation('L = Var(\u2207\u00b2 I_gray)')
add_para('where \u2207\u00b2 is the discrete Laplacian kernel [[0,1,0],[1,-4,1],[0,1,0]].')
add_bullet('High L: Rich edge structure (complex scenes with vegetation, wires, textured backgrounds)')
add_bullet('Low L: Uniform regions (clear sky, smooth surfaces) or severe blur')
add_bullet('Theoretical basis: Captures local intensity curvature; variance measures edge energy distribution (Pech-Pacheco et al., 2000)')

doc.add_heading('Shannon Entropy (H)', level=3)
add_para('The entropy of the grayscale intensity histogram measures information content:')
add_equation('H = \u2212 \u03a3(k=1..B) p_k \u00b7 log\u2082(p_k)')
add_para('where p_k is the normalized frequency of the k-th bin in a B-bin histogram (default B=64).')
add_bullet('High H: Many distinct intensity levels used (complex, diverse scenes)')
add_bullet('Low H: Few intensity levels dominate (uniform scenes, sky)')
add_bullet('Theoretical basis: Shannon (1948) quantifies expected information content; applied to image histograms, measures scene diversity')

doc.add_heading('Tenengrad (Sobel Gradient Energy)', level=3)
add_equation('T = mean(G_x\u00b2 + G_y\u00b2)')
add_para('where G_x, G_y are Sobel gradient responses. Captures directional edge energy, complementary to Laplacian which is isotropic.')

doc.add_heading('NR-IQA Score (BRISQUE-like)', level=3)
add_para('A No-Reference Image Quality Assessment based on Natural Scene Statistics (NSS):')
add_numbered_item('Compute MSCN (Mean Subtracted Contrast Normalized) coefficients:')
add_equation('MSCN(i,j) = (I(i,j) \u2212 \u03bc(i,j)) / (\u03c3(i,j) + \u03b5)')
add_para('where \u03bc and \u03c3 are local Gaussian-weighted mean and standard deviation.')
add_numbered_item('Fit a Generalized Gaussian Distribution (GGD) to MSCN via moment matching:')
add_equation('f(x; \u03b1, \u03b2) = (\u03b2 / (2\u03b1\u0393(1/\u03b2))) \u00b7 exp(\u2212(|x|/\u03b1)^\u03b2)')
add_bullet('Shape parameter \u03b2: higher = more Gaussian (natural), lower = heavier tails (distorted)')
add_numbered_item('Score derived from: GGD shape parameter, shape parameters of paired products, MSCN variance')
add_para('Reference: Mittal et al., "No-Reference Image Quality Assessment in the Spatial Domain," IEEE TIP 2012 (BRISQUE) [2]', italic=True)
add_callout('Note: Our implementation is a simplified BRISQUE-like proxy, NOT standard NIQE. Runs in <3ms on 160\u00d7160 proxy images.', bg_color="E8F8F5", text_color="1E8449")

doc.add_heading('Color Entropy', level=3)
add_para('Joint entropy on H and S channels of HSV color space:')
add_equation('H_color = \u2212 \u03a3 p(h,s) \u00b7 log\u2082(p(h,s))')
add_para('Captures color diversity beyond grayscale information.')

doc.add_heading('2.3 Exponential Moving Average (EMA)', level=2)
add_para('EMA is a recursive low-pass filter that weights recent observations more heavily:')
add_equation('EMA_t = \u03b2 \u00b7 x_t + (1 \u2212 \u03b2) \u00b7 EMA_{t\u22121}')
add_bullet('\u03b2 (smoothing factor): Controls the effective time constant. \u03c4 = 1/\u03b2 frames')
add_bullet('\u03b2=0.30 \u2192 \u03c4~3 frames (fast response); \u03b2=0.02 \u2192 \u03c4~50 frames (slow baseline)')
add_bullet('Dual-EMA: Using fast and slow EMAs together allows detecting signal transitions')

doc.add_heading('2.4 Hysteresis Thresholding', level=2)
add_para('A two-threshold switching mechanism that prevents oscillation (Schmitt trigger analog):')
add_code_block([
    'if current_model == n:',
    '    switch to s if signal >= threshold_high',
    'elif current_model == s:',
    '    switch to n if signal <= threshold_low',
    'else:',
    '    keep current model',
])
add_para('The dead-zone between threshold_low and threshold_high provides switching stability. Essential for video processing where adjacent frames are highly correlated.')

doc.add_heading('2.5 Rolling Percentile Normalization', level=2)
add_para('Maps raw proxy values to [0, 1] using rolling statistics over a sliding window of W frames:')
add_equation('X_norm = clip((X \u2212 P_lo) / (P_hi \u2212 P_lo + \u03b5), 0, 1)')
add_para('where P_lo and P_hi are the 10th and 90th percentiles of the last W values.')
add_callout('\u26a0\ufe0f CENTRAL FINDING: This normalization compresses values to a locally uniform distribution, destroying the absolute magnitude differences that correlate with detection difficulty.', bg_color="FADBD8", text_color="922B21")
doc.add_page_break()

# ==================== CHAPTER 3 ====================
doc.add_heading('3. Related Work', level=1)

doc.add_heading('3.1 UAV-Based Inspection with Deep Learning', level=2)
add_table(
    ['Reference', 'Contribution', 'Relation to DMS'],
    [
        ['Ayoub & Schneider-Kamp (2021) [3]', 'Real-time fault detection on UAVs with onboard DL', 'Single-model; no adaptive switching'],
        ['Xie et al. (2024) [4]', 'YOLOv5s + NVIDIA DeepStream for power-line', 'Fixed model; optimizes but doesn\'t adapt'],
        ['Yang et al. (2024) [5]', 'YOLOv7 + model quantization for damage detection', 'Reduces model size; no dynamic switch'],
        ['Phan et al. (2024) [6]', 'YOLO benchmark for bridge inspection', 'Offline comparison; DMS selects at runtime'],
        ['Ciccone & Ceruti (2025) [7]', 'YOLO for UAV search and rescue', 'Fixed model; no scene-adaptive selection'],
        ['Lyu et al. (2025) [8]', 'Survey: UAV-based DL for civil infrastructure', 'Identifies need for adaptive inference'],
    ]
)

doc.add_heading('3.2 Adaptive and Dynamic Inference', level=2)
add_table(
    ['Reference', 'Contribution', 'Relation to DMS'],
    [
        ['Hu et al. (2019) [9]', 'Anytime prediction with adaptive loss', 'Early-exit; DMS switches between models'],
        ['Sponner et al. (2024) [10]', 'Survey: runtime DNN optimization', 'DMS is "model selection" category'],
        ['EdgeMLBalancer (2025) [11]', 'Self-adaptive model switching on edge', 'No scene complexity or stability'],
        ['MODI (Ogden & Guo, 2019) [12]', 'Mobile deep inference with model selection', 'No scene proxy; no UAV constraints'],
        ['Li et al. (2025) [13]', 'Adaptive switching in multi-CNN UAV swarms', 'Multi-UAV; DMS is single-UAV per-frame'],
        ['Li et al. (2023) [14]', 'Dynamic DNN switching for edge intelligence', 'Collaborative; DMS is standalone'],
        ['Suganya et al. (2024) [15]', 'Dynamic task offloading for UAV edge', 'Offloading; DMS keeps inference on-device'],
    ]
)

doc.add_heading('3.3 Research Gap', level=2)
add_callout('No existing work integrates ALL of the following:', bg_color="D6EAF8", text_color="1B4F72")
add_numbered_item('Scene complexity estimation from image proxies')
add_numbered_item('Per-frame multi-model switching (not just model selection at deployment)')
add_numbered_item('Hysteresis-based switching stability with actuator protections')
add_numbered_item('Detection quality evaluation (not just timing)')
add_numbered_item('Honest proxy failure analysis with negative results')
add_para('DMS-Raptor addresses this gap with a comprehensive 8-policy evaluation framework.', bold=True)
doc.add_page_break()

# ==================== CHAPTER 4 ====================
doc.add_heading('4. System Architecture', level=1)

doc.add_heading('4.1 Pipeline Overview', level=2)
add_para('The DMS-Raptor pipeline processes each video frame through three stages:')
add_numbered_item('Scene Analysis (T_scene: 0\u201330ms) \u2014 Compute image proxies or EMA signals')
add_numbered_item('Controller (T_ctrl: ~0.05\u20130.2ms) \u2014 Normalize, threshold, apply hysteresis and actuators')
add_numbered_item('Inference (T_infer) \u2014 Run ONE selected model: YOLOv8n (~100ms) OR YOLOv8s (~240ms)')
add_code_block([
    'Video Frame',
    '    \u2502',
    '    \u25bc',
    'Scene Analysis (T_scene: 0-30ms)',
    '    \u2502',
    '    \u25bc',
    'Controller (T_ctrl: ~0.05-0.2ms)',
    '    \u2502',
    '    \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510',
    '    \u25bc                      \u25bc',
    'YOLOv8n (~100ms)     YOLOv8s (~240ms)',
    '    \u2502                      \u2502',
    '    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u252c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518',
    '               \u25bc',
    '         Detections + Frame Result',
])

doc.add_heading('4.2 Timing Architecture', level=2)
add_equation('T_total = T_scene + T_ctrl + T_infer   (only ONE model runs per frame)')
add_table(
    ['Component', 'What it measures', 'Typical CPU range'],
    [
        ['T_scene', 'Image proxy computation', '0ms (conf_ema) to 30ms (niqe_switch)'],
        ['T_ctrl', 'Controller logic', '~0.05\u20130.2ms'],
        ['T_infer_n', 'YOLOv8n forward pass + NMS', '~100ms'],
        ['T_infer_s', 'YOLOv8s forward pass + NMS', '~240ms'],
    ]
)
add_callout('Critical: Only ONE model runs per frame. T_total never includes both models.', bg_color="D5F5E3", text_color="1E8449")

doc.add_heading('4.3 Per-Material Model Pairs', level=2)
add_table(
    ['Material', 'n-model', 's-model', 'F1 gap'],
    [
        ['Glass', 'glass_y8n_fast.pt', 'glass_y8s_accurate.pt', '0.870 vs 0.892 (significant)'],
        ['Porcelain', 'porcelain_y8n_fast.pt', 'porcelain_y8s_accurate.pt', '0.962 vs 0.958 (negligible)'],
        ['Composite', 'comp_resplit_y8n_fast.pt', 'comp_resplit_y8s_accurate.pt', '0.957 vs 0.922'],
    ]
)
add_callout('Key Insight: DMS is most beneficial when the model pair has a meaningful accuracy differential (glass >> porcelain).', bg_color="D6EAF8", text_color="1B4F72")
doc.add_page_break()

# ==================== CHAPTER 5 ====================
doc.add_heading('5. Scene Complexity Proxies', level=1)

doc.add_heading('5.1 Proxy Computation Pipeline', level=2)
add_para('All proxies are computed on a downsampled 160\u00d7160 grayscale image for speed:')
add_code_block([
    'img_small = cv2.resize(frame, (160, 160), interpolation=cv2.INTER_AREA)',
    'gray = cv2.cvtColor(img_small, cv2.COLOR_BGR2GRAY)',
])

doc.add_heading('5.2 Individual Proxy Correlations with Model Disagreement', level=2)
add_para('Model disagreement is defined as the set-difference in detections between n and s models (measured via bipartite IoU matching). A good proxy should correlate with disagreement.')
add_table(
    ['Proxy', 'Correlation (glass)', 'Correlation (porcelain)', 'Assessment'],
    [
        ['Raw Laplacian (L)', 'r = \u22120.577 (moderate)', 'r = \u22120.73 (strong)', 'Best image proxy'],
        ['Shannon Entropy (H)', 'r = +0.172 (negligible)', 'r = +0.33 (weak)', 'Contradictory signal'],
        ['Mean Conf(n)', 'r = \u22120.672 (moderate)', '\u2014', 'Best overall proxy'],
        ['NR-IQA Score', 'r = +0.400 (weak)', '\u2014', 'Quality degradation indicator'],
        ['C-score (after norm.)', 'r = \u22120.098 (negligible)', 'r = \u22120.054 (negligible)', 'FAILED'],
    ]
)

doc.add_heading('5.3 Why Raw Proxies Can\'t Be Used Directly', level=2)
add_para('Raw Laplacian values are scene-dependent and non-stationary:')
add_bullet('Video A (rural): L ranges from 50 to 500')
add_bullet('Video B (urban): L ranges from 2000 to 8000')
add_para('A fixed threshold (e.g., "switch if L < 200") would always trigger for Video A and never for Video B. Normalization is needed to make thresholds transferable across videos.')

doc.add_heading('5.4 The Normalization Dilemma', level=2)
add_table(
    ['Method', 'Advantage', 'Disadvantage'],
    [
        ['Rolling percentile', 'Video-agnostic thresholds', 'Destroys absolute magnitude information'],
        ['Fixed global thresholds', 'Preserves signal', 'Doesn\'t generalize across videos'],
        ['EMA-relative', 'Preserves signal dynamics', 'Detects transitions, not absolute complexity'],
        ['Confidence-based (conf_ema)', 'Direct measurement of difficulty', 'One-frame lag; only reactive'],
    ]
)
doc.add_page_break()

# ==================== CHAPTER 6 ====================
doc.add_heading('6. The 8 Switching Policies', level=1)

doc.add_heading('6.1 Baselines', level=2)

doc.add_heading('Policy 1: n_only', level=3)
add_para('Always use YOLOv8n. Fastest possible (T_total ~ 100ms). Lower bound on accuracy.', bold=True)

doc.add_heading('Policy 2: s_only', level=3)
add_para('Always use YOLOv8s. Best accuracy (100% coverage). Upper bound on latency (~240ms). Serves as the pseudo-oracle for detection quality evaluation.', bold=True)

doc.add_heading('6.2 Proactive Policies (decide BEFORE inference)', level=2)

doc.add_heading('Policy 3: entropy_only', level=3)
add_bullet('Signal: Shannon entropy (H) of grayscale histogram')
add_bullet('Decision: Switch to s if H \u2265 median(H), else n')
add_bullet('T_scene: ~3ms (histogram computation)')

doc.add_heading('Policy 4: combined', level=3)
add_bullet('Signal: C-score = \u03b1 \u00b7 L_norm + (1\u2212\u03b1) \u00b7 H_norm (\u03b1=0.6)')
add_bullet('Decision: Switch to s if C \u2265 c_mid (0.50), else n')
add_bullet('T_scene: ~3ms')

doc.add_heading('Policy 5: combined_hyst', level=3)
add_bullet('Signal: Same C-score as combined')
add_bullet('Decision: Hysteresis with c_low=0.45, c_high=0.55')
add_bullet('T_scene: ~3ms')

doc.add_heading('Policy 6: niqe_switch', level=3)
add_bullet('Signal: NR-IQA quality score (BRISQUE-like)')
add_bullet('Decision: Dual-EMA on NIQE score; switch to s when quality degrades')
add_bullet('T_scene: ~3ms (on 160\u00d7160 proxy)')

doc.add_heading('Policy 7: multi_proxy', level=3)
add_bullet('Signal: Weighted composite of 4 proxies with EMA-relative normalization')
add_bullet('Proxies: Laplacian (w=0.37), Tenengrad (w=0.38), Entropy (w=0.29), Color Entropy (w=0.28)')
add_bullet('Normalization: drop_i = |EMA_fast_i \u2212 EMA_slow_i| / (EMA_slow_i + \u03b5)')
add_bullet('Decision: C_multi with hysteresis at c_high=0.10, c_low=0.03')
add_bullet('T_scene: ~8ms (all proxies computed)')

doc.add_heading('6.3 Reactive Policy (uses PREVIOUS frame\'s result)', level=2)

doc.add_heading('Policy 8: conf_ema  \u2b50 Best Speed-Accuracy Trade-off', level=3)
add_bullet('Signal: Dual-EMA on the chosen model\'s detection confidence')
add_bullet('Fast EMA: \u03b2_f = 0.30 (responds in ~3 frames)')
add_bullet('Slow EMA: \u03b2_s = 0.02 (baseline over ~50 frames)')
add_bullet('Rise signal: conf_drop = max(0, (EMA_slow \u2212 EMA_fast) / (EMA_slow + \u03b5))')
add_bullet('Switch to s when conf_drop \u2265 0.12; back to n when conf_drop \u2264 0.04')
add_bullet('T_scene: 0ms (no image processing!)')

add_callout('Why conf_ema is fundamentally different:', bg_color="D5F5E3", text_color="1E8449")
add_numbered_item('Zero overhead: No image processing (T_scene = 0ms)')
add_numbered_item('Direct signal: Observes actual model difficulty, not a proxy for it')
add_numbered_item('Self-correcting: If n-model handles a "complex" scene well, conf_ema stays on n')
add_numbered_item('One-frame lag: Only disadvantage \u2014 uses previous frame\'s confidence')
doc.add_page_break()

# ==================== CHAPTER 7 ====================
doc.add_heading('7. The C-Score Failure (Central Contribution)', level=1)

doc.add_heading('7.1 The Hypothesis', level=2)
add_callout('Hypothesis: Combining Laplacian variance and Shannon entropy via rolling percentile normalization will produce a reliable switching signal.', bg_color="D6EAF8", text_color="1B4F72")
add_para('Intuition: Complex frames have more edges (high L) and more clutter (high H). Normalizing to [0,1] and combining should indicate scene difficulty.')

doc.add_heading('7.2 What Actually Happens', level=2)

add_para('Step 1: Raw proxies have moderate correlation', bold=True, size=12)
add_table(
    ['Proxy', '|r| with Disagreement (glass)', '|r| with Disagreement (porcelain)'],
    [
        ['Raw L', '0.39', '0.73'],
        ['Raw H', '0.17', '0.33'],
    ]
)

add_para('Step 2: Rolling percentile normalization destroys the signal', bold=True, size=12)
add_table(
    ['Signal', '|r| (glass)', '|r| (porcelain)'],
    [
        ['L_norm (after percentile)', '0.20', '0.04'],
        ['C-score (combined normalized)', '0.098', '0.054'],
    ]
)
add_callout('\u26a0\ufe0f Correlation drops from 0.39 \u2192 0.098 (glass) and 0.73 \u2192 0.054 (porcelain)!', bg_color="FADBD8", text_color="922B21")

add_para('Step 3: Why this happens', bold=True, size=12)
add_para('Rolling percentile normalization maps each proxy to a locally uniform distribution within the sliding window:')
add_equation('L_norm = clip((L \u2212 P10) / (P90 \u2212 P10), 0, 1)')
add_para('Within a 200-frame window where L ranges 100\u2013500: L=100\u21920.0, L=500\u21921.0, L=300\u21920.5. But in a different window where L ranges 300\u2013800: L=300 now maps to 0.0 (was 0.5!), L=500 maps to 0.25 (was 1.0!).')
add_callout('The absolute magnitude that correlates with model difficulty is LOST. The normalization tells you where a value sits within the local distribution, not how difficult the scene actually is.', bg_color="FADBD8", text_color="922B21")

doc.add_heading('7.3 Why combined_hyst Still Works (Despite Weak Proxy)', level=2)
add_para('combined_hyst achieves 89% coverage despite near-zero C-score correlation. This is NOT because the proxy is good:')
add_numbered_item('Conservative switching: Hysteresis keeps s-model running for extended periods once activated')
add_numbered_item('High s-usage: Uses s-model for ~50% of frames')
add_numbered_item('Low switching rate: Only 0.83 switches per 100 frames (most stable)')
add_callout('The proxy serves as a random-ish trigger to periodically engage the s-model, not as an accurate classifier.', bg_color="FFF3CD", text_color="856404")

doc.add_heading('7.4 Why This is a Contribution', level=2)
add_numbered_item('Reproducible and well-characterized \u2014 not a vague "it didn\'t work"')
add_numbered_item('Identifies a specific mechanism \u2014 percentile normalization destroying absolute magnitude')
add_numbered_item('Guided the design of conf_ema and multi_proxy \u2014 both avoid rolling percentile normalization')
add_numbered_item('Warns future researchers \u2014 against similar normalization choices in adaptive inference')
doc.add_page_break()

# ==================== CHAPTER 8 ====================
doc.add_heading('8. Actuator Stabilization', level=1)

doc.add_heading('8.1 The Oscillation Problem', level=2)
add_para('Without stabilization, switching policies can oscillate rapidly between n and s models when the signal hovers near the threshold:')
add_bullet('Excessive cache pollution (loading/unloading model weights)')
add_bullet('Unpredictable latency (alternating between 100ms and 240ms)')
add_bullet('No benefit from switching (overhead of instability negates savings)')

doc.add_heading('8.2 Four-Layer Protection', level=2)

add_para('Layer 1: Hysteresis (Dead Zone)', bold=True, size=12, color='2E5D8C')
add_bullet('Switch n\u2192s: signal must exceed c_high')
add_bullet('Switch s\u2192n: signal must drop below c_low')
add_bullet('Dead zone (c_low to c_high): Keep current model')

add_para('Layer 2: Minimum Dwell Time', bold=True, size=12, color='2E5D8C')
add_equation('min_dwell_frames = 10')
add_para('Must stay on new model for at least 10 frames before another switch.')

add_para('Layer 3: Rate Limiting', bold=True, size=12, color='2E5D8C')
add_equation('max_switches_per_100 = 12')
add_para('No more than 12 switches per 100 frames. System locks to current model if limit reached.')

add_para('Layer 4: C-Score EMA Smoothing', bold=True, size=12, color='2E5D8C')
add_equation('C_smoothed = (1 \u2212 \u03b2) \u00b7 C_prev + \u03b2 \u00b7 C_raw    (\u03b2 = 0.25)')
add_para('Smooths the C-score signal before threshold comparison, filtering frame-to-frame noise.')
doc.add_page_break()

# ==================== CHAPTER 9 ====================
doc.add_heading('9. Evaluation Framework', level=1)

doc.add_heading('9.1 Dataset', level=2)
add_table(
    ['Material', 'Videos', 'Total Frames', 'Description'],
    [
        ['Glass', '9', '~65,000', 'Glass insulators on wooden/metal poles'],
        ['Porcelain', '3', '~13,000', 'Porcelain insulators on transmission towers'],
        ['Total', '12', '~78,225', 'UAV inspection footage, various conditions'],
    ]
)

doc.add_heading('9.2 Pseudo-Oracle Validation', level=2)
add_callout('Problem: No ground-truth bounding box annotations exist for these videos.', bg_color="FADBD8", text_color="922B21")
add_para('Solution: Use s_only (YOLOv8s) detections as the pseudo-oracle (ground truth):', bold=True)
add_equation('Detection Coverage = (matched detections) / (total s_only detections)')
add_para('Where "matched" means the policy\'s chosen model produced a detection with IoU \u2265 0.5 matching an s_only detection.')
add_para('Justification: YOLOv8s consistently achieves higher F1 than YOLOv8n on glass (0.892 vs 0.870).', italic=True)

add_para('Metrics computed:', bold=True)
add_bullet('Precision: matched / (matched + false_positives)')
add_bullet('Recall (Coverage): matched / total_oracle_detections')
add_bullet('F1 Score: harmonic mean of precision and recall')
add_bullet('Confidence-Weighted Coverage: \u03a3(conf of matched) / \u03a3(conf of oracle)')

doc.add_heading('9.3 Timing Methodology', level=2)
add_equation('T_total = T_scene + T_ctrl + T_infer')
add_para('Measured with time.perf_counter() (sub-microsecond resolution).')
add_para('NOT included in T_total:', bold=True)
add_bullet('Frame read from disk/camera (I/O)')
add_bullet('Overlay drawing and video writing')
add_bullet('GUI updates or external processing')

doc.add_heading('9.4 Frame-Level Validation', level=2)
add_numbered_item('Run the policy to get routing decisions (n or s per frame)')
add_numbered_item('Sample 30 frames routed to n-model and 30 routed to s-model')
add_numbered_item('Run BOTH models on each sampled frame')
add_numbered_item('Compute IoU between the two models\' detection sets')
add_numbered_item('Compare IoU distributions between n-routed and s-routed frames')
add_callout('Good routing: s-routed frames have LOWER IoU (models disagree \u2192 harder), n-routed have HIGHER IoU (models agree \u2192 easy). IoU gap should be POSITIVE.', bg_color="D5F5E3", text_color="1E8449")
doc.add_page_break()

# ==================== CHAPTER 10 ====================
doc.add_heading('10. Experimental Results (12 Videos, 78,225 frames)', level=1)

add_para('Results below are from 12 insulator inspection videos (9 glass + 3 porcelain) processed on a remote Linux CPU (gchi@134.109.184.66). 3 batch runs. All timing uses StreamingEngine (one model per frame). Detection quality uses s_only as pseudo-oracle. Total: 625,800 policy-frame evaluations.')

doc.add_heading('10.1 Model Baselines (Training Metrics)', level=2)
add_table(
    ['Dataset', 'Model', 'Precision', 'Recall', 'F1', 'mAP50', 'mAP50-95'],
    [
        ['Glass', 'YOLOv8n', '0.881', '0.860', '0.870', '0.902', '0.579'],
        ['Glass', 'YOLOv8s', '0.917', '0.869', '0.892', '0.927', '0.598'],
        ['Porcelain', 'YOLOv8n', '0.957', '0.967', '0.962', '0.977', '0.877'],
        ['Porcelain', 'YOLOv8s', '0.952', '0.964', '0.958', '0.979', '0.882'],
    ]
)

doc.add_heading('10.2 Aggregate Detection Quality (mean across 12 videos)', level=2)
add_table(
    ['Policy', 'Precision (%)', 'Recall (%)', 'F1 (%)', 's-Usage (%)', 'sw/100'],
    [
        ['n_only', '80.8', '78.0', '78.9', '0.0', '0.00'],
        ['s_only (oracle)', '100.0', '100.0', '100.0', '100.0', '0.00'],
        ['entropy_only', '90.1', '87.9', '88.8', '48.2', '1.61'],
        ['combined', '90.1', '87.8', '88.8', '48.6', '1.25'],
        ['combined_hyst', '90.6', '88.2', '89.2', '49.8', '0.87'],
        ['conf_ema \u2b50', '88.8', '87.3', '87.9', '47.9', '4.61'],
        ['niqe_switch', '81.5', '78.9', '79.8', '5.1', '0.08'],
        ['multi_proxy', '83.3', '82.2', '82.4', '20.4', '0.38'],
    ]
)
add_callout('Key: combined_hyst leads on F1 (89.2%), but conf_ema achieves 87.9% at 45% less latency than s_only with zero T_scene overhead. Both are Pareto-optimal.', bg_color="D5F5E3", text_color="1E8449")

doc.add_heading('10.3 Per-Video Detection Quality (12 videos)', level=2)
add_para('Videos ranked by difficulty (n_only F1, ascending):')
add_table(
    ['Video', 'Material', 'Frames', 'n_only F1', 'conf_ema F1', 'Lift', 'comb_hyst F1', 'Lift'],
    [
        ['20190912-233', 'Glass', '9,000', '49.2', '63.5', '+14.3', '67.1', '+17.9'],
        ['UAV_porcelain', 'Porcelain', '1,475', '51.7', '74.5', '+22.8', '70.2', '+18.5'],
        ['glass_ins', 'Glass', '829', '73.4', '85.0', '+11.6', '88.0', '+14.6'],
        ['20190916-722', 'Glass', '8,190', '77.9', '86.4', '+8.5', '90.8', '+12.9'],
        ['20190916-633-01', 'Glass', '9,000', '80.3', '87.5', '+7.2', '91.7', '+11.4'],
        ['161_YUN_0001_96', 'Glass', '8,400', '82.3', '91.3', '+9.0', '89.7', '+7.4'],
        ['20190911-174-01', 'Glass', '9,000', '82.7', '91.4', '+8.7', '93.7', '+11.0'],
        ['porce', 'Porcelain', '4,672', '85.5', '93.7', '+8.2', '95.2', '+9.7'],
        ['porcelain_maybe', 'Porcelain', '5,009', '86.7', '93.0', '+6.3', '94.1', '+7.4'],
        ['20190918-515', 'Glass', '6,480', '86.9', '91.5', '+4.6', '94.0', '+7.1'],
        ['021_YUN_0001_111', 'Glass', '9,000', '94.0', '97.2', '+3.2', '96.8', '+2.8'],
        ['YUN_0001_58', 'Glass', '7,170', '96.7', '99.2', '+2.5', '98.9', '+2.2'],
    ]
)
add_callout('DMS provides the biggest F1 lift on the hardest videos: +22.8% on UAV_porcelain, +17.9% on 20190912-233. On easy videos the lift is +2-3%. DMS value scales with scene difficulty. Porcelain videos confirm cross-material generalization.', bg_color="D6EAF8", text_color="1B4F72")

doc.add_heading('10.4 CPU Timing (12 videos, frame-weighted)', level=2)
add_para('Timing measured with time.perf_counter(). Per-frame timing traces collected across all 78,225 frames. Weighted mean across 12 videos.')
add_table(
    ['Policy', 'T_mean (ms)', 'T_p95 (ms)', 's-route (%)', 'sw/100', 'CV (%)', 'Speedup vs s_only'],
    [
        ['n_only', '78.4', '167.9', '0.0', '0.00', '41.2', '30.9%'],
        ['s_only', '113.4', '176.3', '100.0', '0.00', '27.8', '\u2014'],
        ['entropy_only', '73.9', '98.7', '52.1', '1.61', '23.4', '34.8%'],
        ['combined', '71.6', '97.4', '50.8', '1.25', '20.4', '36.9%'],
        ['combined_hyst', '73.5', '99.2', '51.1', '0.87', '18.9', '35.2%'],
        ['conf_ema \u2b50', '62.4', '94.1', '47.4', '4.61', '11.3', '45.0%'],
        ['niqe_switch', '60.7', '69.5', '3.9', '0.08', '26.7', '46.5%'],
        ['multi_proxy', '63.2', '90.8', '19.2', '0.38', '26.2', '44.3%'],
    ]
)
add_callout('conf_ema delivers 45% latency savings vs s_only with the lowest cross-video CV (11.3%). n_only paradoxically has the worst CV (41.2%) due to outlier spikes.', bg_color="D5F5E3", text_color="1E8449")

doc.add_heading('10.5 The n vs s Latency Gap (12 videos)', level=2)
add_para('The n-s mean latency gap across 12 videos (weighted):')
add_para(f'Baseline gap: n_only = 78.4ms, s_only = 113.4ms, Gap = 35.0ms')
add_para('The gap varies by video resolution and content complexity. Smaller porcelain videos show smaller absolute timing but proportional gaps. The 35ms gap on this server CPU is sufficient for DMS to exploit, providing 45% speedup with conf_ema.')
add_callout('The n-s gap averages 35.0ms on this CPU. On embedded edge platforms (Jetson Nano), the gap would be ~140ms, making DMS benefit proportionally LARGER. DMS is most valuable on the weakest hardware.', bg_color="FFF3CD", text_color="856404")

doc.add_heading('10.6 GPU Results (gchi@134.109.184.66)', level=2)
add_table(
    ['Policy', 'Mean T_total (ms)'],
    [
        ['n_only', '~5.5'],
        ['s_only', '~5.8'],
        ['entropy_only', '~13'],
        ['conf_ema', '~6.0'],
    ]
)
add_callout('\u26a0\ufe0f On GPU, n_only and s_only run at nearly identical speed (~5.5 vs 5.8ms). No latency gap to exploit. DMS is a CPU/edge strategy, NOT a GPU strategy.', bg_color="FADBD8", text_color="922B21")

doc.add_heading('10.7 Cross-Video Consistency (Timing CV, 12 videos)', level=2)
add_para('Coefficient of Variation (CV) of T_total_mean across 12 videos. Lower CV = more predictable performance.')
add_table(
    ['Policy', 'CV (%)', 'Min T_mean (ms)', 'Max T_mean (ms)', 'Median T_mean (ms)', 'Interpretation'],
    [
        ['n_only', '41.2', '36.1', '112.9', '78.5', 'Variable (outlier spikes)'],
        ['s_only', '27.8', '64.0', '143.0', '108.4', 'Variable'],
        ['entropy_only', '23.4', '48.4', '88.3', '74.7', 'Good'],
        ['combined', '20.4', '47.5', '85.9', '72.3', 'Good'],
        ['combined_hyst', '18.9', '49.0', '85.0', '73.3', 'Good'],
        ['conf_ema', '11.3', '53.0', '75.8', '61.4', 'Excellent'],
        ['niqe_switch', '26.7', '34.2', '75.3', '60.1', 'Variable'],
        ['multi_proxy', '26.2', '35.2', '78.8', '64.6', 'Variable'],
    ]
)
add_para('conf_ema has the lowest timing CV (11.3%) \u2014 most predictable latency regardless of input video. n_only paradoxically has the worst CV (41.2%) due to CPU outlier spikes that affect lightweight inference more.')
add_callout('conf_ema is the most DEPLOYABLE policy: predictable latency (CV=11.3%), no scene computation overhead, self-calibrating signal. Combined_hyst is runner-up at CV=18.9%.', bg_color="D5F5E3", text_color="1E8449")

doc.add_heading('10.8 Pareto Front: F1 vs Latency (12 videos)', level=2)
add_para('The Pareto-optimal policies (no other policy is both faster AND more accurate):')
add_bullet('n_only: F1=78.9%, T_mean=78.4ms \u2014 fast baseline, lowest accuracy')
add_bullet('conf_ema: F1=87.9%, T_mean=62.4ms \u2014 Pareto-optimal (best speed-accuracy trade-off, 45% faster than s_only)')
add_bullet('combined_hyst: F1=89.2%, T_mean=73.5ms \u2014 Pareto-optimal (best DMS accuracy, most stable switching)')
add_bullet('s_only: F1=100%, T_mean=113.4ms \u2014 perfect accuracy, slowest')
add_para('Non-Pareto: niqe_switch (fast but essentially n_only \u2014 DEAD), multi_proxy (under-switches at 19.2% s-usage), entropy_only (good but dominated by combined_hyst).', italic=True)
doc.add_page_break()

# ==================== CHAPTER 11 ====================
doc.add_heading('11. Key Research Findings', level=1)

doc.add_heading('Finding 1: C-Score Normalization Failure (Central Contribution)', level=2)
add_para('Rolling percentile normalization destroys proxy-target correlation. Raw Laplacian has |r|=0.39 (glass), but after normalization C-score drops to |r|=0.098. This is a fundamental mismatch between normalization and switching objective, not a tuning problem.')

doc.add_heading('Finding 2: combined_hyst Works Through Conservatism, Not Accuracy', level=2)
add_para('Despite near-zero proxy correlation, combined_hyst achieves F1=89.2% \u2014 highest among DMS policies across 12 videos. It uses s-model for 51.1% of frames with only 0.87 switches/100 (most stable). The proxy serves as a random-ish trigger to periodically engage the s-model, not as an accurate scene classifier.')

doc.add_heading('Finding 3: conf_ema is the Pareto-Optimal Speed-Accuracy Trade-off', level=2)
add_para('F1=87.9% at 62.4ms mean (45% faster than s_only). Zero image processing cost (T_scene=0ms). Only 47.4% s-model usage. Lowest cross-video CV=11.3% \u2014 most predictable. The only disadvantage is slightly higher switching rate (4.61/100) due to reactive nature.')

doc.add_heading('Finding 4: DMS Benefit Scales with Video Difficulty', level=2)
add_para('On the hardest videos (20190912-233 F1=49.2%, UAV_porcelain F1=51.7%), DMS provides +14\u201323% F1 lift. On the easiest (YUN_0001_58, n_only F1=96.7%), the lift is only +2.5%. DMS is most valuable when the n-model struggles. This holds across BOTH glass and porcelain insulator types.')

doc.add_heading('Finding 5: GPU Deployment Boundary', level=2)
add_para('On GPU, n_only \u2248 s_only (~5.5 vs 5.8ms). No latency gap to exploit. DMS is designed for CPU/edge deployment where the 28\u201358ms gap (measured) or 100\u2013140ms gap (embedded platforms) exists.')

doc.add_heading('Finding 6: niqe_switch and multi_proxy Are Ineffective', level=2)
add_para('niqe_switch uses s-model for only 3.9% of frames across 12 videos (F1=79.8%, barely above n_only at 78.9%). The NR-IQA signal is too stable frame-to-frame to trigger meaningful switching. multi_proxy at 19.2% s-usage achieves F1=82.4% \u2014 better than niqe_switch but well below the ~50% target split needed for optimal DMS performance.')

doc.add_heading('Finding 7: The n-s Latency Gap Determines DMS Viability', level=2)
add_para('The measured n-s gap on the test CPU is 35.0ms (n=78.4ms, s=113.4ms). DMS provides 45% speedup via conf_ema. The gap is smaller than on embedded edge platforms (Jetson Nano: ~140ms), meaning DMS benefit would be proportionally LARGER on real edge devices. DMS is a CPU/edge strategy; on GPU the gap is ~0ms (both models ~18-22ms).')

doc.add_heading('Finding 8: Cross-Material Generalization (NEW)', level=2)
add_para('With the addition of 3 porcelain videos in batch 3, DMS demonstrates consistent behavior across insulator types. Porcelain videos show the same switching patterns, similar s-usage splits, and comparable F1 improvements. The hardest porcelain video (UAV_porcelain, n_F1=51.7%) receives the largest DMS lift (+22.8%), consistent with the difficulty-scaling pattern seen in glass videos.')

doc.add_page_break()

# ==================== CHAPTER 12 ====================
doc.add_heading('12. Implementation Details', level=1)

doc.add_heading('12.1 Single-File Architecture', level=2)
add_para('The entire experiment pipeline is implemented in dms_experiment.py (~1700 lines):')
add_code_block([
    'dms_experiment.py',
    '\u251c\u2500\u2500 Configuration (RunConfig, InferenceParams, FrameResult, RunSummary)',
    '\u251c\u2500\u2500 NR-IQA Score (compute_nriqa_score)',
    '\u251c\u2500\u2500 Image Proxies (complexity_proxies_fast, extended proxies)',
    '\u251c\u2500\u2500 StreamingEngine (all 8 policies in one engine)',
    '\u251c\u2500\u2500 PolicySimulator (replay policy logic on precomputed signals)',
    '\u251c\u2500\u2500 Detection Quality Analysis (run_detection_quality)',
    '\u251c\u2500\u2500 Frame-Level Validation (run_frame_validation)',
    '\u251c\u2500\u2500 Overnight Pipeline (run_overnight_pipeline)',
    '\u251c\u2500\u2500 Plot Generation (generate_all_plots, 15+ plot functions)',
    '\u251c\u2500\u2500 Report Generator (generate_report)',
    '\u2514\u2500\u2500 CLI (argparse with subcommands)',
])

doc.add_heading('12.2 Key Design Decisions', level=2)
add_numbered_item('No secondary NMS: YOLO\'s .predict() applies NMS internally. An earlier bug added O(n\u00b2) Python NMS on top, making n SLOWER than s (340ms vs 333ms). Removing it restored correct ratio.')
add_numbered_item('Per-material models: Glass and porcelain have different visual characteristics. Separate model pairs outperform composite.')
add_numbered_item('conf_ema cross-pipeline fix: In detection quality mode, EMA update must use confidence from the model the policy would have chosen, not always n-model.')
add_numbered_item('Proxy computation on 160\u00d7160: Downsampling from 1080p reduces proxy time from ~50ms to ~3ms with negligible information loss.')

doc.add_heading('12.3 Experiment Workflow', level=2)
add_code_block([
    'Local Machine (Windows)              Remote (Linux, CPU)',
    '  \u2502                                     \u2502',
    '  \u251c\u2500\u2500 dms_experiment.py  \u2500\u2500push\u2500\u2500>     \u2502',
    '  \u2502                                     \u251c\u2500\u2500 Download 1 video from GDrive',
    '  \u2502                                     \u251c\u2500\u2500 Detection quality',
    '  \u2502                                     \u251c\u2500\u2500 Frame validation',
    '  \u2502                                     \u251c\u2500\u2500 Overnight pipeline (CPU)',
    '  \u2502                                     \u251c\u2500\u2500 Upload results to GDrive',
    '  \u2502                                     \u251c\u2500\u2500 Delete video + results',
    '  \u2502                                     \u2514\u2500\u2500 Repeat for next video',
    '  \u2502                                     \u2502',
    '  \u2502<\u2500\u2500 download from GDrive \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518',
    '  \u251c\u2500\u2500 Generate plots locally',
    '  \u2514\u2500\u2500 Generate thesis report',
])
doc.add_page_break()

# ==================== CHAPTER 13 ====================
doc.add_heading('13. Deployment Considerations', level=1)


doc.add_heading('13.1 When to Use DMS', level=2)
add_table(
    ['Condition', 'DMS Benefit', 'Recommended Policy'],
    [
        ['CPU/edge with n/s gap > 50ms', 'High', 'conf_ema or combined_hyst'],
        ['High-end GPU (n \u2248 s latency)', 'None', 'Just use s_only'],
        ['Model pair with F1 gap > 0.02', 'High', 'conf_ema'],
        ['Negligible F1 gap', 'Low', 'Just use n_only'],
        ['Latency-critical', 'Medium', 'conf_ema (zero T_scene)'],
        ['Accuracy-critical', 'Medium', 'combined_hyst (89% coverage)'],
    ]
)

doc.add_heading('13.2 Edge Platform Optimization', level=2)
add_para('For NVIDIA Jetson Nano deployment:')
add_bullet('TensorRT FP16/INT8 quantization: 3\u20135x inference speedup [16]')
add_bullet('cv2.cuda.resize: GPU-accelerated proxy computation')
add_bullet('Hardware video decoder: Frame read without CPU overhead [17]')
add_bullet('Shared memory: Model weights loaded once, reducing memory pressure')

doc.add_heading('13.3 Limitations', level=2)
add_numbered_item('No ground truth: Pseudo-oracle validation assumes s_only is correct')
add_numbered_item('One-frame lag (conf_ema): Reactive approach can miss first frame of difficulty spike')
add_numbered_item('Per-material training: Requires separate model pairs per insulator material')
add_numbered_item('No multi-class support: Current evaluation is single-class (insulator detection only)')
doc.add_page_break()

# ==================== CHAPTER 14 ====================
doc.add_heading('14. Future Work: Data & Model Agnosticism', level=1)

add_para('The current DMS-Raptor evaluation is limited to UAV insulator inspection with YOLOv8n/s model pairs. To establish DMS as a general-purpose framework, validation across different datasets, object domains, and model architectures is essential.')

doc.add_heading('14.1 The Agnosticism Hypothesis', level=2)
add_para('DMS switching policies are fundamentally model-agnostic and domain-agnostic \u2014 they only require:', bold=True)
add_numbered_item('A fast model (lower latency, lower accuracy)')
add_numbered_item('An accurate model (higher latency, higher accuracy)')
add_numbered_item('A measurable latency gap between them on the deployment platform')
add_para('The switching signal (confidence EMA, entropy, etc.) depends only on model output statistics, not on the specific object class or model architecture.')

doc.add_heading('14.2 Proposed Validation Experiments', level=2)

add_para('Dataset Agnosticism \u2014 Different object domains:', bold=True, size=12, color='2E5D8C')
add_table(
    ['Dataset', 'Domain', 'Source', 'Why'],
    [
        ['COCO val2017', 'General objects (80 classes)', 'Pre-trained YOLOv8', 'Standard benchmark, no training needed'],
        ['VisDrone', 'Aerial/drone detection', 'Pre-trained or fine-tuned', 'Closest to UAV inspection domain'],
        ['BDD100K', 'Autonomous driving', 'Pre-trained', 'Different camera angle, urban scenes'],
        ['DOTA', 'Aerial object detection', 'Fine-tuned', 'Overhead view, small objects'],
    ]
)

add_para('Model Agnosticism \u2014 Different model pairs:', bold=True, size=12, color='2E5D8C')
add_table(
    ['Model Pair', 'Fast Model', 'Accurate Model', 'Expected Gap (CPU)'],
    [
        ['YOLOv8 n/s', 'YOLOv8n (3.2M)', 'YOLOv8s (11.2M)', '~35ms (measured)'],
        ['YOLOv8 n/m', 'YOLOv8n (3.2M)', 'YOLOv8m (25.9M)', '~80\u2013120ms'],
        ['YOLOv8 s/l', 'YOLOv8s (11.2M)', 'YOLOv8l (43.7M)', '~100\u2013200ms'],
        ['YOLOv9 t/s', 'YOLOv9t', 'YOLOv9s', 'Architecture-level test'],
        ['YOLOv11 n/s', 'YOLOv11n', 'YOLOv11s', 'Latest generation'],
        ['RT-DETR l/x', 'RT-DETR-l', 'RT-DETR-x', 'Transformer-based (non-YOLO)'],
    ]
)

doc.add_heading('14.3 Experimental Design', level=2)
add_para('For each dataset \u00d7 model pair combination:')
add_numbered_item('Run DMS with all 8 policies (same parameters as insulator experiments)')
add_numbered_item('Measure: F1 vs oracle, timing (median, p95), s-usage %, switching rate')
add_numbered_item('Compare conf_ema vs baselines (n_only, s_only)')
add_numbered_item('Record: CPU timing AND GPU timing (to verify deployment boundary)')
add_numbered_item('Use actual ground-truth annotations (not pseudo-oracle) where available')

add_para('Expected Outcome:', bold=True)
add_bullet('conf_ema should remain Pareto-optimal across domains (it uses model confidence, not domain-specific features)')
add_bullet('C-score failure should persist (normalization is domain-independent)')
add_bullet('GPU deployment boundary should hold (latency gap \u2248 0 on GPU regardless of model)')
add_bullet('DMS benefit should scale with the n-s accuracy gap (larger gap = more room for intelligent switching)')

doc.add_heading('14.4 Reinforcement Learning Extension', level=2)
add_para('Replace hand-tuned threshold policies with learned switching agents:')
add_bullet('LinUCB contextual bandit: 12-dim feature vector (image proxies + temporal signals), learns optimal switching boundary')
add_bullet('Thompson Sampling: Bayesian posterior over switching actions, natural exploration-exploitation balance')
add_bullet('Key question: Does LinUCB learn to weight features differently than the hand-crafted C-score?')
add_callout('If LinUCB feature importance shows that detection confidence dominates (matching conf_ema), it validates the reactive approach. If it finds image proxies useful, it suggests the hand-tuned thresholds were suboptimal.', bg_color="D6EAF8", text_color="1B4F72")

doc.add_heading('14.5 Comprehensive Results Table (Target)', level=2)
add_para('The final agnosticism study should produce a table of this form:')
add_table(
    ['Dataset', 'Model Pair', 'Platform', 'n-s Gap', 'conf_ema F1', 'Speedup', 'DMS Viable?'],
    [
        ['Insulators (glass)', 'YOLOv8 n/s', 'CPU', '35ms', '86.7%', '32.1%', 'Yes'],
        ['Insulators (porcelain)', 'YOLOv8 n/s', 'CPU', 'TBD', 'TBD', 'TBD', 'TBD'],
        ['COCO val2017', 'YOLOv8 n/s', 'CPU', 'TBD', 'TBD', 'TBD', 'TBD'],
        ['COCO val2017', 'YOLOv8 n/m', 'CPU', 'TBD', 'TBD', 'TBD', 'TBD'],
        ['COCO val2017', 'YOLOv8 n/s', 'GPU', 'TBD', 'TBD', 'TBD', 'No (expected)'],
        ['VisDrone', 'YOLOv8 n/s', 'CPU', 'TBD', 'TBD', 'TBD', 'TBD'],
        ['VisDrone', 'YOLOv11 n/s', 'CPU', 'TBD', 'TBD', 'TBD', 'TBD'],
    ]
)
doc.add_page_break()

# ==================== CHAPTER 15 ====================
doc.add_heading('15. References', level=1)

refs = [
    '[1]  G. Jocher, A. Chaurasia, and J. Qiu, "Ultralytics YOLOv8," 2023. github.com/ultralytics/ultralytics',
    '[2]  A. Mittal, A. K. Moorthy, and A. C. Bovik, "No-Reference Image Quality Assessment in the Spatial Domain," IEEE Trans. Image Process., vol. 21, no. 12, pp. 4695\u20134708, 2012.',
    '[3]  N. Ayoub and P. Schneider-Kamp, "Real-time on-board deep learning fault detection for autonomous UAV inspections," Electronics, vol. 10, no. 9, p. 1091, 2021.',
    '[4]  S. Xie et al., "Real-time object detection from UAV inspection videos by combining YOLOv5s and DeepStream," Sensors, vol. 24, no. 12, p. 3862, 2024.',
    '[5]  X. Yang et al., "UAV-deployed deep learning network for real-time multi-damage detection using model quantization techniques," Autom. Constr., vol. 159, p. 105254, 2024.',
    '[6]  T.-N. Phan et al., "Deep learning models for UAV-assisted bridge inspection: A YOLO benchmark analysis," arXiv:2411.04475, 2024.',
    '[7]  F. Ciccone and A. Ceruti, "Real-time search and rescue with drones: A deep learning approach for small-object detection based on YOLO," Drones, vol. 9, no. 8, p. 514, 2025.',
    '[8]  C. Lyu et al., "UAV-based deep learning applications for automated inspection of civil infrastructure," Autom. Constr., vol. 177, p. 106285, 2025.',
    '[9]  H. Hu et al., "Learning anytime predictions in neural networks via adaptive loss balancing," Proc. AAAI, vol. 33, pp. 3812\u20133821, 2019.',
    '[10] M. Sponner et al., "Adapting neural networks at runtime: Current trends in at-runtime optimizations for deep learning," ACM Comput. Surv., vol. 56, no. 10, art. 248, 2024.',
    '[11] A. Matathammal et al., "EdgeMLBalancer: A self-adaptive approach for dynamic model switching on resource-constrained edge devices," arXiv:2502.06493, 2025.',
    '[12] S. S. Ogden and T. Guo, "MODI: Mobile deep inference made efficient by edge computing," Proc. USENIX HotEdge, 2019.',
    '[13] Y. Li et al., "Adaptive model switching of collaborative inference for multi-CNN streams in UAV swarm," Chinese J. Aeronaut., vol. 38, no. 8, p. 103564, 2025.',
    '[14] Y. Li et al., "Dynamic DNN model switching for collaborative edge intelligence in UAV swarm," Proc. WCSP, pp. 923\u2013929, 2023.',
    '[15] B. Suganya et al., "Dynamic task offloading edge-aware optimization framework for enhanced UAV operations on edge computing platform," Sci. Rep., vol. 14, p. 16383, 2024.',
    '[16] NVIDIA, "DeepStream SDK developer guide," 2023. developer.nvidia.com/deepstream-sdk',
    '[17] B. Battseren, "Software architecture for real-time image analysis in autonomous MAV missions," Ph.D. dissertation, TU Chemnitz, 2024.',
    '[18] Z. Jiang et al., "Efficient deep learning inference on edge devices," Proc. SysML Conf., 2018.',
    '[19] I. J. Ratul et al., "Accelerating deep learning inference: A comparative analysis of modern acceleration frameworks," Electronics, vol. 14, no. 15, p. 2977, 2025.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-1.5)
    p.paragraph_format.space_after = Pt(3)
    # Split reference number from text
    if ref.startswith('['):
        bracket_end = ref.index(']') + 1
        run_num = p.add_run(ref[:bracket_end])
        run_num.bold = True
        run_num.font.size = Pt(10)
        run_num.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        run_text = p.add_run(ref[bracket_end:])
        run_text.font.size = Pt(10)
    else:
        run = p.add_run(ref)
        run.font.size = Pt(10)

# ---- Footer note ----
doc.add_paragraph()
add_separator()
add_para('This document serves as a complete theoretical reference for the DMS-Raptor thesis. '
         'All findings, numbers, and methodology are based on the experimental pipeline implemented in dms_experiment.py.',
         italic=True, size=10, color='888888', align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- Save ----
out_path = r'G:\Teja_Master_Thesis\DMS_Clean\DMS_Thesis_Theory_Handbook_v2.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
